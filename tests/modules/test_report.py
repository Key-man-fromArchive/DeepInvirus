# @TASK B1-B9 - Word report generation tests (updated for v2 framework)
# @SPEC docs/planning/10-workplan-v2-report-framework.md#Phase-B
# @TEST tests/modules/test_report.py
"""Tests for bin/generate_report.py - universal virome report framework.

Covers:
- Valid .docx file generation from mock input data
- Report structure validation (v2 section headings: Executive Summary through Limitations)
- Figure insertion into the Word document
- Table insertion into the Word document
- CLI argument parsing
- B1: New report structure (0-9 + Appendix)
- B2: Methods auto-generation (minimap2, scipy)
- B3: Scientific hedging in conclusions
- B4: VIRUS_ORIGIN dict structure
- B5: Top virus auto-detection
- B6: Conditional diversity section
- B7: Limitations auto-generation
- B8: QC waterfall table
- B9: Universal family descriptions
"""

from __future__ import annotations

import subprocess
import sys
import textwrap
from pathlib import Path

import numpy as np
import pandas as pd
import pytest
from docx import Document

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parents[2]
BIN_DIR = PROJECT_ROOT / "bin"
GENERATE_REPORT_SCRIPT = BIN_DIR / "generate_report.py"


# ---------------------------------------------------------------------------
# Fixtures: mock input data
# ---------------------------------------------------------------------------


@pytest.fixture
def mock_bigtable_tsv(tmp_path: Path) -> Path:
    """Create a mock bigtable.tsv file."""
    df = pd.DataFrame(
        {
            "seq_id": [f"seq_{i}" for i in range(12)],
            "sample": ["S1"] * 4 + ["S2"] * 4 + ["S3"] * 4,
            "seq_type": ["contig"] * 12,
            "length": [5000, 3500, 4200, 2800, 5100, 3400, 4100, 2900, 4800, 3600, 4300, 3000],
            "detection_method": ["genomad", "diamond", "both", "genomad"] * 3,
            "detection_score": [0.95, 0.87, 0.92, 0.78, 0.93, 0.85, 0.91, 0.80, 0.94, 0.86, 0.90, 0.79],
            "taxid": [10001, 10002, 10003, 10004] * 3,
            "domain": ["Viruses"] * 12,
            "phylum": ["Nucleocytoviricota", "Leviviricota", "Nucleocytoviricota", "Pisuviricota"] * 3,
            "class": ["Megaviricetes", "Leviviricetes", "Megaviricetes", "Pisoniviricetes"] * 3,
            "order": ["Imitervirales", "Norzivirales", "Imitervirales", "Picornavirales"] * 3,
            "family": ["Iflaviridae", "Narnaviridae", "Baculoviridae", "Dicistroviridae"] * 3,
            "genus": ["Iflavirus", "Narnavirus", "Alphabaculovirus", "Cripavirus"] * 3,
            "species": ["Iflavirus sp.", "Narnavirus sp.", "Alphabaculovirus sp.", "Cripavirus sp."] * 3,
            "ictv_classification": ["ICTV_A", "ICTV_B", "ICTV_C", "ICTV_D"] * 3,
            "baltimore_group": ["ssRNA(+)", "ssRNA(+)", "dsDNA", "ssRNA(+)"] * 3,
            "count": [150, 120, 130, 90, 160, 110, 140, 95, 155, 115, 135, 85],
            "rpm": [150.5, 120.3, 130.2, 90.1, 160.0, 110.5, 140.3, 95.2, 155.1, 115.4, 135.0, 85.3],
            "coverage": [25.5, 12.3, 18.2, 9.1, 28.0, 11.5, 17.3, 8.2, 23.1, 13.4, 16.0, 7.3],
        }
    )
    path = tmp_path / "bigtable.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


@pytest.fixture
def mock_matrix_tsv(tmp_path: Path) -> Path:
    """Create a mock sample_taxon_matrix.tsv file."""
    df = pd.DataFrame(
        {
            "taxon": ["Iflaviridae", "Narnaviridae", "Baculoviridae", "Dicistroviridae"],
            "taxid": [10001, 10002, 10003, 10004],
            "rank": ["family"] * 4,
            "S1": [150.5, 120.3, 130.2, 90.1],
            "S2": [160.0, 110.5, 140.3, 95.2],
            "S3": [155.1, 115.4, 135.0, 85.3],
        }
    )
    path = tmp_path / "sample_taxon_matrix.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


@pytest.fixture
def mock_alpha_tsv(tmp_path: Path) -> Path:
    """Create a mock alpha_diversity.tsv file."""
    df = pd.DataFrame(
        {
            "sample": ["S1", "S2", "S3"],
            "observed_species": [4, 4, 4],
            "shannon": [1.35, 1.32, 1.34],
            "simpson": [0.74, 0.73, 0.74],
            "chao1": [4.0, 4.0, 4.0],
            "pielou_evenness": [0.97, 0.95, 0.96],
        }
    )
    path = tmp_path / "alpha_diversity.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


@pytest.fixture
def mock_pcoa_tsv(tmp_path: Path) -> Path:
    """Create a mock pcoa_coordinates.tsv file."""
    df = pd.DataFrame(
        {
            "sample": ["S1", "S2", "S3"],
            "PC1": [0.25, -0.15, -0.10],
            "PC2": [0.10, -0.20, 0.10],
            "PC3": [0.05, 0.03, -0.08],
        }
    )
    path = tmp_path / "pcoa_coordinates.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


@pytest.fixture
def mock_qc_stats_tsv(tmp_path: Path) -> Path:
    """Create a mock fastp QC stats file."""
    df = pd.DataFrame(
        {
            "sample": ["S1", "S2", "S3"],
            "raw_reads": [10000000, 12000000, 11000000],
            "raw_bases": [1500000000, 1800000000, 1650000000],
            "trimmed_reads": [9500000, 11400000, 10450000],
            "trimmed_bases": [1425000000, 1710000000, 1567500000],
            "q30_rate": [0.95, 0.94, 0.95],
            "gc_content": [0.42, 0.43, 0.42],
            "host_removed_reads": [8000000, 9600000, 8800000],
        }
    )
    path = tmp_path / "qc_stats.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


@pytest.fixture
def mock_assembly_stats_tsv(tmp_path: Path) -> Path:
    """Create a mock assembly stats file."""
    df = pd.DataFrame(
        {
            "sample": ["S1", "S2", "S3"],
            "total_contigs": [1500, 1800, 1650],
            "total_length": [3000000, 3600000, 3300000],
            "n50": [5000, 4500, 4800],
            "largest_contig": [25000, 22000, 23000],
            "gc_percent": [42.1, 43.2, 42.5],
        }
    )
    path = tmp_path / "assembly_stats.tsv"
    df.to_csv(path, sep="\t", index=False)
    return path


# ---------------------------------------------------------------------------
# Helper: run report generation
# ---------------------------------------------------------------------------

def _run_report(tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
                mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
                extra_args=None):
    """Run generate_report.py and return (result, output_path)."""
    output = tmp_path / "report.docx"
    cmd = [
        sys.executable,
        str(GENERATE_REPORT_SCRIPT),
        "--bigtable", str(mock_bigtable_tsv),
        "--matrix", str(mock_matrix_tsv),
        "--alpha", str(mock_alpha_tsv),
        "--pcoa", str(mock_pcoa_tsv),
        "--qc-stats", str(mock_qc_stats_tsv),
        "--assembly-stats", str(mock_assembly_stats_tsv),
        "--output", str(output),
    ]
    if extra_args:
        cmd.extend(extra_args)
    result = subprocess.run(cmd, capture_output=True, text=True)
    return result, output


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestGenerateReportScript:
    """Test the generate_report.py script exists and is importable."""

    def test_script_exists(self) -> None:
        """generate_report.py must exist in bin/."""
        assert GENERATE_REPORT_SCRIPT.exists(), (
            f"Script not found: {GENERATE_REPORT_SCRIPT}"
        )

    def test_script_importable(self) -> None:
        """The script module must be importable without errors."""
        result = subprocess.run(
            [sys.executable, "-c",
             f"import importlib.util; spec = importlib.util.spec_from_file_location('generate_report', '{GENERATE_REPORT_SCRIPT}'); mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod)"],
            capture_output=True,
            text=True,
        )
        assert result.returncode == 0, f"Import failed: {result.stderr}"


@pytest.mark.unit
class TestDocxGeneration:
    """Test that a valid .docx file is generated with correct v2 structure."""

    def test_generates_valid_docx(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Running generate_report.py must produce a valid .docx file."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        assert result.returncode == 0, f"Script failed:\nSTDOUT: {result.stdout}\nSTDERR: {result.stderr}"
        assert output.exists(), "Output .docx file was not created"
        assert output.stat().st_size > 0, "Output .docx file is empty"

        doc = Document(str(output))
        assert doc is not None, "Failed to open .docx with python-docx"

    def test_report_v2_section_headings(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Report must contain all v2 section headings (B1 redesign)."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))

        headings = []
        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                headings.append(para.text)

        heading_text = " | ".join(headings)

        # B1: Expected v2 sections
        expected_sections = [
            "Executive Summary",
            "Methods",
            "QC Results",
            "Host Removal",
            "Virus Detection",
            "Coverage Analysis",
            "Taxonomic Analysis",
            "Diversity",
            "Conclusions",
            "Limitations",
            "Appendix",
        ]

        for section in expected_sections:
            found = any(section in h for h in headings)
            assert found, (
                f"Section heading containing '{section}' not found. "
                f"Found headings: {heading_text}"
            )

    def test_report_contains_tables(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Report must contain tables (project info, detection, diversity, etc.)."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        assert len(doc.tables) >= 3, (
            f"Expected at least 3 tables in report, found {len(doc.tables)}"
        )

    def test_report_contains_figures(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Report must contain inline images (figures)."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))

        image_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                inline_shapes = run._r.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
                )
                image_count += len(inline_shapes)

        assert image_count >= 2, (
            f"Expected at least 2 figures in report, found {image_count}"
        )

    def test_figures_directory_created(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Figures should also be saved to a figures/ directory."""
        figures_dir = tmp_path / "figures"
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
            extra_args=["--figures-dir", str(figures_dir)],
        )
        assert result.returncode == 0, f"Script failed: {result.stderr}"
        assert figures_dir.exists(), "figures/ directory not created"

        png_files = list(figures_dir.glob("*.png"))
        assert len(png_files) >= 2, (
            f"Expected at least 2 PNG figures, found {len(png_files)}"
        )


@pytest.mark.unit
class TestReportAppendix:
    """Test that the report appendix sections are present."""

    def test_appendix_present(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Report must contain appendix section."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs
                     if p.style and p.style.name and p.style.name.startswith("Heading")]

        found_appendix = any("Appendix" in h for h in headings)
        assert found_appendix, (
            f"Appendix heading not found. Headings: {headings}"
        )


# ---------------------------------------------------------------------------
# B2: Methods auto-generation tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestMethodsAutoGeneration:
    """B2: Verify methods text uses correct tool names (no hardcoding errors)."""

    def test_no_bowtie2_in_methods(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Methods must NOT reference Bowtie2 (replaced by minimap2)."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)

        assert "Bowtie2" not in all_text, (
            "Bowtie2 should not appear in the report (replaced by minimap2)"
        )

    def test_no_scikit_bio_in_methods(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Methods must NOT reference scikit-bio (replaced by scipy + numpy)."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)

        assert "scikit-bio" not in all_text, (
            "scikit-bio should not appear in the report (replaced by scipy + numpy)"
        )

    def test_minimap2_present(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Methods must mention minimap2."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "minimap2" in all_text, "minimap2 must be mentioned in the report"

    def test_scipy_numpy_present(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Methods must mention scipy + numpy."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "scipy" in all_text, "scipy must be mentioned in the report"
        assert "numpy" in all_text, "numpy must be mentioned in the report"


# ---------------------------------------------------------------------------
# B3: Scientific hedging tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestScientificHedging:
    """B3: Conclusions must use hedged language, not definitive statements."""

    def test_no_dead_alive_language(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Report must NOT use 'dead sample' / 'alive sample' language."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)

        forbidden = [
            "죽은 샘플", "살아있는 샘플", "세포 사멸로 인해",
            "dead sample", "alive sample", "cell death caused",
        ]
        for phrase in forbidden:
            assert phrase not in all_text, (
                f"Forbidden phrase '{phrase}' found in report (B3: no causal attribution)"
            )

    def test_no_active_replication_claim(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Report must NOT claim 'active viral replication' from coverage alone."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)

        forbidden = ["활발한 바이러스 증식", "active viral replication"]
        for phrase in forbidden:
            assert phrase not in all_text, (
                f"Forbidden phrase '{phrase}' found (B3: coverage != replication)"
            )


# ---------------------------------------------------------------------------
# B4: VIRUS_ORIGIN structure tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestVirusOrigin:
    """B4: VIRUS_ORIGIN dictionary structure validation."""

    def test_virus_origin_importable(self) -> None:
        """VIRUS_ORIGIN and VIRUS_ORIGIN_CLASS_FALLBACK must be importable."""
        result = subprocess.run(
            [sys.executable, "-c",
             "import importlib.util; "
             f"spec = importlib.util.spec_from_file_location('gr', '{GENERATE_REPORT_SCRIPT}'); "
             "mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod); "
             "assert hasattr(mod, 'VIRUS_ORIGIN'); "
             "assert hasattr(mod, 'VIRUS_ORIGIN_CLASS_FALLBACK'); "
             "assert 'Iflaviridae' in mod.VIRUS_ORIGIN; "
             "assert 'Picornaviridae' not in mod.VIRUS_ORIGIN; "
             "assert 'Caudoviricetes' in mod.VIRUS_ORIGIN_CLASS_FALLBACK"],
            capture_output=True,
            text=True,
        )
        assert result.returncode == 0, f"VIRUS_ORIGIN validation failed: {result.stderr}"

    def test_virus_origin_has_confidence(self) -> None:
        """Every VIRUS_ORIGIN entry must have 'origin' and 'confidence' keys."""
        result = subprocess.run(
            [sys.executable, "-c",
             "import importlib.util; "
             f"spec = importlib.util.spec_from_file_location('gr', '{GENERATE_REPORT_SCRIPT}'); "
             "mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod); "
             "errors = [f for f, v in mod.VIRUS_ORIGIN.items() if 'origin' not in v or 'confidence' not in v]; "
             "assert not errors, f'Missing keys in: {errors}'"],
            capture_output=True,
            text=True,
        )
        assert result.returncode == 0, f"VIRUS_ORIGIN key check failed: {result.stderr}"


# ---------------------------------------------------------------------------
# B5: Top virus auto-detection tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestTopVirusDetection:
    """B5: detect_top_virus must work correctly."""

    def test_detect_top_virus_returns_series(self) -> None:
        """detect_top_virus should return a Series for valid data."""
        result = subprocess.run(
            [sys.executable, "-c",
             "import importlib.util, pandas as pd; "
             f"spec = importlib.util.spec_from_file_location('gr', '{GENERATE_REPORT_SCRIPT}'); "
             "mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod); "
             "bt = pd.DataFrame({'seq_id': ['s1','s2'], 'family': ['Iflaviridae','Narnaviridae'], "
             "'length': [5000, 3000], 'coverage': [25.0, 10.0]}); "
             "top = mod.detect_top_virus(bt); "
             "assert top is not None; "
             "assert top['family'] == 'Iflaviridae'"],
            capture_output=True,
            text=True,
        )
        assert result.returncode == 0, f"detect_top_virus failed: {result.stderr}"

    def test_detect_top_virus_no_parvo_hardcode(self) -> None:
        """detect_top_virus must NOT hardcode Parvoviridae."""
        result = subprocess.run(
            [sys.executable, "-c",
             "import importlib.util, inspect; "
             f"spec = importlib.util.spec_from_file_location('gr', '{GENERATE_REPORT_SCRIPT}'); "
             "mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod); "
             "src = inspect.getsource(mod.detect_top_virus); "
             "assert 'Parvoviridae' not in src, 'detect_top_virus must not hardcode Parvoviridae'"],
            capture_output=True,
            text=True,
        )
        assert result.returncode == 0, f"Parvoviridae hardcode check failed: {result.stderr}"


# ---------------------------------------------------------------------------
# B7: Limitations auto-generation tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestLimitationsGeneration:
    """B7: Limitations section must be auto-generated based on context."""

    def test_limitations_present_in_report(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Report must contain a Limitations section."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs
                     if p.style and p.style.name and p.style.name.startswith("Heading")]

        found = any("Limitations" in h for h in headings)
        assert found, f"Limitations heading not found. Headings: {headings}"

    def test_limitations_mentions_rnaseq(
        self, tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
        mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
    ) -> None:
        """Limitations must mention RNA-seq DNA virus caveat."""
        result, output = _run_report(
            tmp_path, mock_bigtable_tsv, mock_matrix_tsv, mock_alpha_tsv,
            mock_pcoa_tsv, mock_qc_stats_tsv, mock_assembly_stats_tsv,
        )
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "DNA virus" in all_text or "DNA viruses" in all_text, (
            "Limitations must mention RNA-seq DNA virus caveat"
        )


# ---------------------------------------------------------------------------
# B8: QC waterfall tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestQCWaterfall:
    """B8: QC waterfall table builder."""

    def test_qc_waterfall_function(self) -> None:
        """_build_qc_waterfall must produce a DataFrame with expected columns."""
        result = subprocess.run(
            [sys.executable, "-c",
             "import importlib.util, pandas as pd; "
             f"spec = importlib.util.spec_from_file_location('gr', '{GENERATE_REPORT_SCRIPT}'); "
             "mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod); "
             "bbduk = [{'sample': 'S1', 'total_reads': 10000, 'adapter_removed': 500, "
             "'adapter_pct': 5.0, 'phix_removed': 10, 'clean_reads': 9500}]; "
             "host = pd.DataFrame({'sample': ['S1'], 'unmapped_reads': [5000]}); "
             "wf = mod._build_qc_waterfall(bbduk, host); "
             "assert not wf.empty; "
             "assert 'Sample' in wf.columns; "
             "assert 'Raw Reads' in wf.columns; "
             "assert 'After Adapter' in wf.columns"],
            capture_output=True,
            text=True,
        )
        assert result.returncode == 0, f"QC waterfall test failed: {result.stderr}"


# ---------------------------------------------------------------------------
# B9: Family descriptions universalization tests
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestFamilyDescriptions:
    """B9: Family descriptions must be universal (no insect-specific language)."""

    def test_no_insect_specific_language(self) -> None:
        """FAMILY_DESCRIPTIONS must not contain insect-specific interpretation phrases."""
        script = (
            "import importlib.util, sys\n"
            f"spec = importlib.util.spec_from_file_location('gr', '{GENERATE_REPORT_SCRIPT}')\n"
            "mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod)\n"
            "forbidden = ['곤충에서의 검출은', '곤충 장내', '곤충 시료에서의', '곤충이 식물을', '곤충 장내 진균']\n"
            "errors = []\n"
            "for fam, desc in mod.FAMILY_DESCRIPTIONS.items():\n"
            "    for phrase in forbidden:\n"
            "        if phrase in desc:\n"
            "            errors.append(f'{fam}: {phrase!r}')\n"
            "if errors:\n"
            "    print('FAIL: ' + '; '.join(errors), file=sys.stderr)\n"
            "    sys.exit(1)\n"
        )
        result = subprocess.run(
            [sys.executable, "-c", script],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, f"Insect-specific language check failed: {result.stderr}"

    def test_no_picornaviridae_crpv_dcv(self) -> None:
        """If Picornaviridae is in FAMILY_DESCRIPTIONS, it must not mention CrPV/DCV."""
        script = (
            "import importlib.util, sys\n"
            f"spec = importlib.util.spec_from_file_location('gr', '{GENERATE_REPORT_SCRIPT}')\n"
            "mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod)\n"
            "if 'Picornaviridae' in mod.FAMILY_DESCRIPTIONS:\n"
            "    desc = mod.FAMILY_DESCRIPTIONS['Picornaviridae']\n"
            "    assert 'CrPV' not in desc, 'CrPV belongs to Dicistroviridae'\n"
            "    assert 'DCV' not in desc, 'DCV belongs to Dicistroviridae'\n"
            "    assert 'Cricket paralysis' not in desc\n"
        )
        result = subprocess.run(
            [sys.executable, "-c", script],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, f"Picornaviridae CrPV/DCV check failed: {result.stderr}"
