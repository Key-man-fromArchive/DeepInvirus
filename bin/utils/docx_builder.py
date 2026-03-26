# @TASK T0.7 - Word report builder utility
# @SPEC docs/planning/05-design-system.md#4-타이포그래피
# @SPEC docs/planning/05-design-system.md#5-word-보고서-템플릿
"""Word (.docx) report builder for DeepInvirus analysis reports.

Wraps python-docx to provide a simple API for generating formatted
Word documents with headings, paragraphs, tables, and figures that
comply with the DeepInvirus design-system specification
(05-design-system.md section 4.2 & 5.2).
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# @TASK T0.7 - Design-system Word typography (section 4.2)
# ---------------------------------------------------------------------------
_FONT_NAME = "Malgun Gothic"  # 맑은 고딕
_TITLE_SIZE = Pt(20)
_SUBTITLE_SIZE = Pt(14)
_BODY_SIZE = Pt(11)
_TABLE_SIZE = Pt(9)
_CAPTION_SIZE = Pt(10)

# Page setup (section 5.2)
_PAGE_WIDTH = Cm(21.0)  # A4
_PAGE_HEIGHT = Cm(29.7)
_MARGIN_TOP = Cm(2.54)
_MARGIN_BOTTOM = Cm(2.54)
_MARGIN_LEFT = Cm(3.17)
_MARGIN_RIGHT = Cm(3.17)

# Table styling
_HEADER_BG = RGBColor(0x1F, 0x77, 0xB4)  # Deep Blue from palette
_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
_ALT_ROW_BG = RGBColor(0xF8, 0xF9, 0xFA)  # Surface from palette


class ReportBuilder:
    """Builder for Word (.docx) analysis reports.

    Provides a fluent interface for constructing reports that follow
    the DeepInvirus design-system specification.

    Example::

        builder = ReportBuilder()
        builder.add_heading("Analysis Overview", level=1)
        builder.add_paragraph("This report summarises ...")
        builder.add_table(summary_df, title="Table 1. Summary")
        builder.add_figure(Path("figures/heatmap.png"), "Figure 1.")
        builder.save(Path("report.docx"))
    """

    def __init__(self, template_path: Optional[Path] = None) -> None:
        """Initialise the report builder.

        Args:
            template_path: Optional path to a .docx template file.
                If ``None``, a new blank document is created with
                default DeepInvirus page settings applied.

        Raises:
            FileNotFoundError: If the template file does not exist.
        """
        if template_path is not None:
            template_path = Path(template_path)
            if not template_path.exists():
                raise FileNotFoundError(
                    f"Template not found: {template_path}"
                )
            self._doc = Document(str(template_path))
            logger.info("Loaded template from %s", template_path)
        else:
            self._doc = Document()
            logger.info("Created new blank document")

        self._apply_page_setup()

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _apply_page_setup(self) -> None:
        """Apply DeepInvirus page-layout settings (section 5.2)."""
        for section in self._doc.sections:
            section.page_width = _PAGE_WIDTH
            section.page_height = _PAGE_HEIGHT
            section.orientation = WD_ORIENT.PORTRAIT
            section.top_margin = _MARGIN_TOP
            section.bottom_margin = _MARGIN_BOTTOM
            section.left_margin = _MARGIN_LEFT
            section.right_margin = _MARGIN_RIGHT

            # Header text
            header = section.header
            header.is_linked_to_previous = False
            if not header.paragraphs:
                header.add_paragraph()
            hp = header.paragraphs[0]
            hp.text = "DeepInvirus Analysis Report"
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in hp.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

            # Footer with page number
            footer = section.footer
            footer.is_linked_to_previous = False
            if not footer.paragraphs:
                footer.add_paragraph()
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Insert PAGE field
            run = fp.add_run()
            fld_char_begin = run._r.makeelement(qn("w:fldChar"), {})
            fld_char_begin.set(qn("w:fldCharType"), "begin")
            run._r.append(fld_char_begin)

            run2 = fp.add_run()
            instr = run2._r.makeelement(qn("w:instrText"), {})
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run2._r.append(instr)

            run3 = fp.add_run()
            fld_char_end = run3._r.makeelement(qn("w:fldChar"), {})
            fld_char_end.set(qn("w:fldCharType"), "end")
            run3._r.append(fld_char_end)

    def _set_run_font(
        self,
        run: "docx.text.run.Run",
        size: Pt = _BODY_SIZE,
        bold: bool = False,
        italic: bool = False,
        color: Optional[RGBColor] = None,
    ) -> None:
        """Apply font settings to a single run."""
        run.font.name = _FONT_NAME
        run.font.size = size
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

        # Ensure East-Asian font is also set (for Korean glyphs)
        rpr = run._r.get_or_add_rPr()
        ea_font = rpr.makeelement(qn("w:rFonts"), {})
        ea_font.set(qn("w:eastAsia"), _FONT_NAME)
        rpr.insert(0, ea_font)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def add_table_of_contents(self, title: str = "Table of Contents") -> None:
        """Insert a Table of Contents field that Word will populate on open.

        The TOC uses built-in Heading styles (Heading 1-3) and creates
        clickable hyperlinks to each heading in the document.

        Note:
            The TOC field is rendered when the document is opened in Word
            and the user presses F9 or selects 'Update Field'.  It will
            show 'Right-click to update field' until then.
        """
        # Title for TOC page
        heading = self._doc.add_heading(title, level=1)
        for run in heading.runs:
            self._set_run_font(run, size=_TITLE_SIZE, bold=True)

        para = self._doc.add_paragraph()
        run = para.add_run()

        # Begin TOC field
        fld_begin = run._r.makeelement(qn("w:fldChar"), {})
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        # TOC instruction: \o "1-3" = heading levels 1-3, \h = hyperlinks
        run2 = para.add_run()
        instr = run2._r.makeelement(qn("w:instrText"), {})
        instr.set(qn("xml:space"), "preserve")
        instr.text = r' TOC \o "1-3" \h \z \u '
        run2._r.append(instr)

        # Separate field code
        run3 = para.add_run()
        fld_sep = run3._r.makeelement(qn("w:fldChar"), {})
        fld_sep.set(qn("w:fldCharType"), "separate")
        run3._r.append(fld_sep)

        # Placeholder text
        run4 = para.add_run("Right-click and select 'Update Field' to generate table of contents")
        self._set_run_font(run4, size=_BODY_SIZE, italic=True,
                          color=RGBColor(0x7F, 0x7F, 0x7F))

        # End TOC field
        run5 = para.add_run()
        fld_end = run5._r.makeelement(qn("w:fldChar"), {})
        fld_end.set(qn("w:fldCharType"), "end")
        run5._r.append(fld_end)

        # Page break after TOC
        self._doc.add_page_break()
        logger.info("Added Table of Contents field")

    def add_heading(self, text: str, level: int = 1) -> None:
        """Add a heading to the report.

        Args:
            text: Heading text.
            level: Heading level (1 = top-level, 2 = subsection, etc.).
                Clamped to 1..4.

        Raises:
            ValueError: If level is not between 1 and 4.
        """
        if not 1 <= level <= 4:
            raise ValueError(f"Heading level must be 1-4, got {level}")

        heading = self._doc.add_heading(text, level=level)
        size = _TITLE_SIZE if level == 1 else _SUBTITLE_SIZE
        for run in heading.runs:
            self._set_run_font(run, size=size, bold=True)

    def add_paragraph(
        self, text: str, style: str = "Normal"
    ) -> None:
        """Add a body paragraph to the report.

        Args:
            text: Paragraph text.
            style: Word paragraph style name (default ``"Normal"``).
        """
        para = self._doc.add_paragraph(style=style)
        run = para.add_run(text)
        self._set_run_font(run, size=_BODY_SIZE)

        # Line spacing 1.5 (section 5.2)
        para_format = para.paragraph_format
        para_format.line_spacing = 1.5

    def add_table(self, df: pd.DataFrame, title: str = "") -> None:
        """Add a DataFrame as a formatted Word table.

        Args:
            df: Data to render.  The index is *not* included; use
                ``df.reset_index()`` beforehand if needed.
            title: Optional title rendered as a bold paragraph above
                the table.

        Raises:
            ValueError: If the DataFrame is empty.
        """
        if df.empty:
            raise ValueError("Cannot add empty DataFrame as table")

        if title:
            title_para = self._doc.add_paragraph()
            run = title_para.add_run(title)
            self._set_run_font(run, size=_BODY_SIZE, bold=True)

        n_rows, n_cols = df.shape
        table = self._doc.add_table(
            rows=n_rows + 1, cols=n_cols
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row
        for j, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col_name)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._set_run_font(
                        run,
                        size=_TABLE_SIZE,
                        bold=True,
                        color=_HEADER_FG,
                    )
            # Header background
            shading = cell._tc.get_or_add_tcPr().makeelement(
                qn("w:shd"), {}
            )
            shading.set(qn("w:fill"), _HEADER_BG.__str__().replace("#", ""))
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for i in range(n_rows):
            for j in range(n_cols):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(df.iloc[i, j])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._set_run_font(run, size=_TABLE_SIZE)

                # Alternating row background
                if i % 2 == 1:
                    shading = cell._tc.get_or_add_tcPr().makeelement(
                        qn("w:shd"), {}
                    )
                    shading.set(
                        qn("w:fill"),
                        _ALT_ROW_BG.__str__().replace("#", ""),
                    )
                    shading.set(qn("w:val"), "clear")
                    cell._tc.get_or_add_tcPr().append(shading)

        logger.info("Added table '%s' (%d rows x %d cols)", title, n_rows, n_cols)

    def add_figure(
        self,
        image_path: Path,
        caption: str = "",
        width_inches: float = 6.0,
    ) -> None:
        """Insert an image with an optional caption.

        Args:
            image_path: Path to the image file (PNG, JPEG, SVG, etc.).
            caption: Figure caption rendered below the image in italic.
            width_inches: Display width in inches (default 6.0).

        Raises:
            FileNotFoundError: If the image file does not exist.
        """
        image_path = Path(image_path)
        if not image_path.exists():
            raise FileNotFoundError(f"Image not found: {image_path}")

        # Centre-aligned image paragraph
        para = self._doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))

        # Caption
        if caption:
            cap_para = self._doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            self._set_run_font(
                cap_run, size=_CAPTION_SIZE, italic=True
            )

        logger.info("Added figure: %s", image_path.name)

    def save(self, output_path: Path) -> Path:
        """Save the document to disk.

        Args:
            output_path: Destination file path (.docx).

        Returns:
            The resolved output path.

        Raises:
            OSError: If the file cannot be written.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(output_path))
        logger.info("Report saved to %s", output_path)
        return output_path
